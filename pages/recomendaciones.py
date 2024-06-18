import streamlit as st
import io
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
from utils.data_manager import extraer_texto_pdf
import openai

def show():
    st.title("✍️ IA para Modificación de Documentos")

    uploaded_file = st.file_uploader("Sube un documento (PDF o DOCX)", type=['pdf', 'docx'])

    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            texto_original = extraer_texto_pdf(uploaded_file)
            st.text_area("Contenido del Documento Original", texto_original, height=300)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            documento = Document(uploaded_file)
            texto_original = "\n".join([para.text for para in documento.paragraphs])
            st.text_area("Contenido del Documento Original", texto_original, height=300)

        st.markdown("### Modificar Documento")

        receptor = st.text_input("Modificar Para (Nombre del Receptor):")
        notas_adicionales = st.text_area("Notas Adicionales:", height=100)

        if st.button("Generar Documento Modificado"):
            contenido_modificado = modificar_documento(texto_original, receptor, notas_adicionales)
            st.text_area("Contenido del Documento Modificado", contenido_modificado, height=300)

            # Descargar documento modificado
            if uploaded_file.type == "application/pdf":
                pdf_modificado = modificar_pdf(uploaded_file, contenido_modificado)
                st.download_button(
                    label="📥 Descargar PDF Modificado",
                    data=pdf_modificado,
                    file_name="documento_modificado.pdf",
                    mime="application/pdf"
                )
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                docx_modificado = modificar_docx(contenido_modificado)
                st.download_button(
                    label="📥 Descargar DOCX Modificado",
                    data=docx_modificado,
                    file_name="documento_modificado.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

def modificar_documento(texto, receptor, notas):
    """Modifica el contenido del documento basado en el receptor y notas adicionales."""
    prompt = f"Modifica el siguiente texto para que esté dirigido a {receptor} y añade las siguientes notas: {notas}.\n\nTexto:\n{texto}"
    response = openai.Completion.create(
        engine="gpt-3.5-turbo",
        prompt=prompt,
        max_tokens=1500
    )
    return response.choices[0].text.strip()

def modificar_pdf(pdf_file, nuevo_contenido):
    """Modifica el contenido de un PDF."""
    reader = PdfReader(pdf_file)
    writer = PdfWriter()

    # Añadir el contenido modificado a una nueva página del PDF
    for page in reader.pages:
        writer.add_page(page)

    nueva_pagina = writer.add_blank_page(width=reader.pages[0].mediabox.width, height=reader.pages[0].mediabox.height)
    nueva_pagina.insert_text(nuevo_contenido)

    pdf_output = io.BytesIO()
    writer.write(pdf_output)
    pdf_output.seek(0)

    return pdf_output

def modificar_docx(nuevo_contenido):
    """Modifica el contenido de un DOCX."""
    doc = Document()
    for line in nuevo_contenido.split("\n"):
        doc.add_paragraph(line)

    docx_output = io.BytesIO()
    doc.save(docx_output)
    docx_output.seek(0)

    return docx_output
